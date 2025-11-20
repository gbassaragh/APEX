"""
Unit tests for DocumentParser Excel/Word parsing.
"""
from io import BytesIO

import pytest

from apex.services.document_parser import DocumentParser


@pytest.fixture
def parser():
    return DocumentParser()


@pytest.mark.asyncio
async def test_parse_excel_extracts_sheets(parser):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Test Sheet"
    ws["A1"] = "Header 1"
    ws["B1"] = "Header 2"
    ws["A2"] = "Value 1"
    ws["B2"] = "Value 2"

    excel_bytes = BytesIO()
    wb.save(excel_bytes)
    excel_bytes.seek(0)

    result = await parser._parse_excel(excel_bytes.read(), "test.xlsx")

    assert result["filename"] == "test.xlsx"
    assert len(result["sheets"]) == 1
    assert result["sheets"][0]["name"] == "Test Sheet"
    assert result["sheets"][0]["rows"][1][0] == "Value 1"


@pytest.mark.asyncio
async def test_parse_excel_extracts_metadata(parser):
    """Test Excel metadata extraction (workbook properties)."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data Sheet"
    ws["A1"] = "Test"

    excel_bytes = BytesIO()
    wb.save(excel_bytes)
    excel_bytes.seek(0)

    result = await parser._parse_excel(excel_bytes.read(), "test.xlsx")

    assert result["metadata"]["format"] == "excel"
    assert result["metadata"]["sheet_count"] == 1
    assert "workbook_properties" in result["metadata"]


@pytest.mark.asyncio
async def test_parse_excel_skips_empty_rows(parser):
    """Test that Excel parsing skips completely empty rows."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Row 1"
    # Row 2 is empty
    ws["A3"] = "Row 3"

    excel_bytes = BytesIO()
    wb.save(excel_bytes)
    excel_bytes.seek(0)

    result = await parser._parse_excel(excel_bytes.read(), "test.xlsx")

    # Should only have 2 rows (empty row skipped)
    assert len(result["sheets"][0]["rows"]) == 2


@pytest.mark.asyncio
async def test_parse_excel_error_handling(parser):
    """Test Excel parsing error handling with corrupted file."""
    corrupted_bytes = b"Not an Excel file"

    with pytest.raises(ValueError, match="Failed to parse Excel file"):
        await parser._parse_excel(corrupted_bytes, "corrupted.xlsx")


@pytest.mark.asyncio
async def test_parse_word_extracts_paragraphs(parser):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Test paragraph 1")
    doc.add_paragraph("Test paragraph 2")

    word_bytes = BytesIO()
    doc.save(word_bytes)
    word_bytes.seek(0)

    result = await parser._parse_word(word_bytes.read(), "test.docx")

    assert result["filename"] == "test.docx"
    assert len(result["paragraphs"]) == 2
    assert result["paragraphs"][0]["text"] == "Test paragraph 1"


@pytest.mark.asyncio
async def test_parse_word_extracts_tables(parser):
    """Test Word table extraction."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Introduction")

    # Add table with 2 rows, 3 columns
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(0, 2).text = "Header 3"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    table.cell(1, 2).text = "Data 3"

    word_bytes = BytesIO()
    doc.save(word_bytes)
    word_bytes.seek(0)

    result = await parser._parse_word(word_bytes.read(), "test.docx")

    assert len(result["tables"]) == 1
    assert result["tables"][0]["row_count"] == 2
    assert result["tables"][0]["column_count"] == 3
    assert result["tables"][0]["cells"][0][0] == "Header 1"
    assert result["tables"][0]["cells"][1][2] == "Data 3"


@pytest.mark.asyncio
async def test_parse_word_extracts_metadata(parser):
    """Test Word metadata extraction."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    doc.add_table(rows=1, cols=2)

    word_bytes = BytesIO()
    doc.save(word_bytes)
    word_bytes.seek(0)

    result = await parser._parse_word(word_bytes.read(), "test.docx")

    assert result["metadata"]["format"] == "word"
    assert result["metadata"]["paragraph_count"] >= 2
    assert result["metadata"]["table_count"] == 1


@pytest.mark.asyncio
async def test_parse_word_error_handling(parser):
    """Test Word parsing error handling with corrupted file."""
    corrupted_bytes = b"Not a Word document"

    with pytest.raises(ValueError, match="Failed to parse Word document"):
        await parser._parse_word(corrupted_bytes, "corrupted.docx")
