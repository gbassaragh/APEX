"""
Document parsing service with Azure Document Intelligence integration.

CRITICAL: PDFs MUST use Azure Document Intelligence, not PyPDF2.
Circuit breaker pattern for Azure DI resilience.
Dead letter queue for failed documents.
"""
import asyncio
import logging
from datetime import datetime
from enum import Enum
from typing import Any, Dict, Optional

from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
from azure.core.exceptions import HttpResponseError

from apex.azure.auth import get_azure_credential
from apex.azure.blob_storage import BlobStorageClient
from apex.config import config
from apex.utils.errors import BusinessRuleViolation

logger = logging.getLogger(__name__)


class CircuitState(Enum):
    """Circuit breaker states for Azure Document Intelligence."""

    CLOSED = "closed"  # Normal operation
    OPEN = "open"  # Failing, reject requests
    HALF_OPEN = "half_open"  # Testing recovery


class CircuitBreaker:
    """
    Circuit breaker for Azure Document Intelligence operations.

    Prevents cascading failures by temporarily disabling calls after
    threshold failures are reached.

    States:
    - CLOSED: Normal operation, calls allowed
    - OPEN: Failing, reject all calls (return error immediately)
    - HALF_OPEN: Testing recovery (allow ONE probe request only)

    Thresholds from config:
    - AZURE_DI_CIRCUIT_BREAKER_THRESHOLD: Failures before opening
    - AZURE_DI_CIRCUIT_BREAKER_TIMEOUT: Seconds before retry

    Thread Safety:
    - Uses asyncio.Lock to prevent race conditions in HALF_OPEN state
    - Only one probe request allowed when testing recovery
    """

    def __init__(self):
        """Initialize circuit breaker in CLOSED state."""
        self.state = CircuitState.CLOSED
        self.failure_count = 0
        self.last_failure_time: Optional[datetime] = None
        self.threshold = config.AZURE_DI_CIRCUIT_BREAKER_THRESHOLD
        self.timeout_seconds = config.AZURE_DI_CIRCUIT_BREAKER_TIMEOUT
        self._lock = asyncio.Lock()  # Prevent HALF_OPEN race conditions
        self._half_open_probe_active = False  # Only one probe at a time

    async def record_success(self) -> None:
        """Record successful operation, reset failure count."""
        async with self._lock:
            if self.state == CircuitState.HALF_OPEN:
                logger.info("Circuit breaker: Probe succeeded, closing circuit")
                self.state = CircuitState.CLOSED
                self._half_open_probe_active = False

            self.failure_count = 0
            self.last_failure_time = None

    async def record_failure(self) -> None:
        """Record failed operation, potentially open circuit."""
        async with self._lock:
            self.failure_count += 1
            self.last_failure_time = datetime.utcnow()

            # Transition HALF_OPEN → OPEN on probe failure
            if self.state == CircuitState.HALF_OPEN:
                logger.error("Circuit breaker: Probe failed, reopening circuit")
                self.state = CircuitState.OPEN
                self._half_open_probe_active = False

            # Transition CLOSED → OPEN on threshold breach
            elif self.failure_count >= self.threshold and self.state == CircuitState.CLOSED:
                logger.error(
                    f"Circuit breaker: Opening circuit after {self.failure_count} failures. "
                    f"Will retry in {self.timeout_seconds}s"
                )
                self.state = CircuitState.OPEN

    async def can_proceed(self) -> bool:
        """
        Check if operation can proceed.

        Returns:
            True if operation allowed

        Raises:
            BusinessRuleViolation: If circuit is open (service unavailable)
                                   or HALF_OPEN probe already in progress
        """
        async with self._lock:
            if self.state == CircuitState.CLOSED:
                return True

            if self.state == CircuitState.OPEN:
                # Check if timeout has elapsed
                if self.last_failure_time:
                    elapsed = (datetime.utcnow() - self.last_failure_time).total_seconds()
                    if elapsed >= self.timeout_seconds:
                        logger.info(
                            f"Circuit breaker: Timeout elapsed ({elapsed:.1f}s), entering HALF_OPEN"
                        )
                        self.state = CircuitState.HALF_OPEN
                        self._half_open_probe_active = True
                        return True

                # Still in timeout period
                elapsed = (datetime.utcnow() - self.last_failure_time).total_seconds()
                remaining = self.timeout_seconds - elapsed
                raise BusinessRuleViolation(
                    message=(
                        f"Azure Document Intelligence circuit breaker is OPEN. "
                        f"Service temporarily unavailable. Try again in {remaining:.0f}s"
                    ),
                    code="CIRCUIT_BREAKER_OPEN",
                )

            # HALF_OPEN - allow only ONE probe request
            if self.state == CircuitState.HALF_OPEN:
                if self._half_open_probe_active:
                    # Probe already in progress - reject this request
                    raise BusinessRuleViolation(
                        message="Circuit breaker probe in progress, please retry shortly",
                        code="CIRCUIT_BREAKER_PROBE_ACTIVE",
                    )
                else:
                    # This is the probe request
                    self._half_open_probe_active = True
                    return True

            return True


class DocumentParser:
    """
    Document parsing service with Azure Document Intelligence integration.

    Features:
    - Azure Document Intelligence for PDFs (mandatory)
    - openpyxl for Excel files
    - python-docx for Word documents
    - Circuit breaker pattern for resilience
    - Dead letter queue for failed documents
    - Async polling with configurable timeout

    Example:
        ```python
        parser = DocumentParser()
        structured_data = await parser.parse_document(
            document_bytes=pdf_bytes,
            filename="scope.pdf",
            blob_path="uploads/project123/scope.pdf"
        )
        ```
    """

    def __init__(self):
        """
        Initialize document parser.

        Client initialization is deferred until first use.
        """
        self._di_client: Optional[DocumentIntelligenceClient] = None
        self._blob_client: Optional[BlobStorageClient] = None
        self._circuit_breaker = CircuitBreaker()

    async def _get_di_client(self) -> DocumentIntelligenceClient:
        """
        Get or create Azure Document Intelligence client.

        Lazy initialization pattern.

        Returns:
            DocumentIntelligenceClient instance
        """
        if self._di_client is None:
            credential = await get_azure_credential()
            self._di_client = DocumentIntelligenceClient(
                endpoint=config.AZURE_DI_ENDPOINT, credential=credential
            )
            logger.info(f"Initialized DocumentIntelligenceClient for {config.AZURE_DI_ENDPOINT}")

        return self._di_client

    async def _get_blob_client(self) -> BlobStorageClient:
        """Get or create blob storage client (async for consistency)."""
        if self._blob_client is None:
            self._blob_client = BlobStorageClient()

        return self._blob_client

    async def parse_document(
        self, document_bytes: bytes, filename: str, blob_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse document and extract structured content.

        Routes to appropriate parser based on file extension:
        - .pdf → Azure Document Intelligence (mandatory)
        - .xlsx, .xls → openpyxl
        - .docx, .doc → python-docx

        Args:
            document_bytes: Binary document content
            filename: Original filename (used for extension detection)
            blob_path: Optional blob path for DLQ (format: "container/blob_name")

        Returns:
            Structured document data with pages, tables, paragraphs

        Raises:
            BusinessRuleViolation: If circuit breaker is open or unsupported format
            TimeoutError: If parsing exceeds timeout
            HttpResponseError: Azure Document Intelligence errors

        Example:
            ```python
            result = await parser.parse_document(
                document_bytes=pdf_bytes,
                filename="scope.pdf",
                blob_path="uploads/project123/scope.pdf"
            )
            # result = {
            #     "pages": [...],
            #     "tables": [...],
            #     "paragraphs": [...],
            #     "metadata": {...}
            # }
            ```
        """
        file_ext = filename.lower().split(".")[-1]

        logger.info(f"Parsing document: {filename} ({len(document_bytes)} bytes, type: {file_ext})")

        try:
            if file_ext == "pdf":
                return await self._parse_pdf_with_azure_di(document_bytes, filename, blob_path)
            elif file_ext in ["xlsx", "xls"]:
                return await self._parse_excel(document_bytes, filename)
            elif file_ext in ["docx", "doc"]:
                return await self._parse_word(document_bytes, filename)
            else:
                # Unsupported format - don't DLQ, this is user error not document failure
                raise BusinessRuleViolation(
                    message=f"Unsupported file format: .{file_ext}. Supported: PDF, Excel, Word",
                    code="UNSUPPORTED_FORMAT",
                )

        except BusinessRuleViolation:
            # Circuit breaker or unsupported format - don't DLQ, bubble up to caller
            # These are not document-specific failures
            raise

        except (HttpResponseError, TimeoutError) as exc:
            # Document parsing failures - DLQ eligible
            logger.error(f"Document parsing failed for {filename}: {exc}")

            if blob_path:
                await self._handle_parsing_failure(blob_path, filename, exc)

            raise

        except Exception as exc:
            # Unexpected errors - log and bubble up (don't DLQ unknown errors)
            logger.error(f"Unexpected error parsing {filename}: {exc}", exc_info=True)
            raise

    async def _parse_pdf_with_azure_di(
        self, document_bytes: bytes, filename: str, blob_path: Optional[str]
    ) -> Dict[str, Any]:
        """
        Parse PDF using Azure Document Intelligence.

        Implements circuit breaker pattern and async polling with timeout.

        Args:
            document_bytes: PDF binary content
            filename: Original filename
            blob_path: Optional blob path for DLQ

        Returns:
            Structured PDF content

        Raises:
            BusinessRuleViolation: If circuit breaker is open
            TimeoutError: If analysis exceeds timeout
            HttpResponseError: Azure DI errors
        """
        # Check circuit breaker (await - it's async now)
        await self._circuit_breaker.can_proceed()

        try:
            client = await self._get_di_client()

            logger.info(f"Starting Azure DI analysis for {filename}")

            # Begin analysis with prebuilt-layout model
            # Use document= parameter with content_type for raw bytes
            poller = await client.begin_analyze_document(
                model_id="prebuilt-layout", document=document_bytes, content_type="application/pdf"
            )

            # Use SDK's built-in polling with timeout (replaces custom _poll_with_timeout)
            try:
                result = await poller.result(timeout=config.AZURE_DI_TIMEOUT)
            except asyncio.TimeoutError:
                # SDK raises asyncio.TimeoutError on timeout
                timeout = config.AZURE_DI_TIMEOUT
                msg = (
                    f"Azure Document Intelligence analysis exceeded "
                    f"{timeout}s timeout for {filename}"
                )
                raise TimeoutError(msg)

            # Extract structured content
            structured_data = self._extract_structured_content(result, filename)

            # Record success in circuit breaker (await - it's async now)
            await self._circuit_breaker.record_success()

            logger.info(f"Successfully parsed PDF: {filename}")

            return structured_data

        except HttpResponseError as exc:
            # Record failure in circuit breaker (await - it's async now)
            await self._circuit_breaker.record_failure()

            logger.error(
                f"Azure DI analysis failed for {filename}: {exc.status_code} {exc.message}"
            )
            raise

        except TimeoutError:
            # Record failure in circuit breaker (await - it's async now)
            await self._circuit_breaker.record_failure()
            raise

    def _extract_structured_content(self, result: Any, filename: str) -> Dict[str, Any]:
        """
        Extract structured content from Azure DI result.

        Args:
            result: Azure DI analysis result
            filename: Document filename

        Returns:
            Structured data with pages, tables, paragraphs

        Note:
            This is a simplified extraction. Production implementation should
            preserve table row/column structure and paragraph hierarchy.
        """
        structured = {
            "filename": filename,
            "pages": [],
            "tables": [],
            "paragraphs": [],
            "metadata": {
                "page_count": len(result.pages) if hasattr(result, "pages") else 0,
                "model_id": "prebuilt-layout",
                "api_version": result.api_version if hasattr(result, "api_version") else None,
            },
        }

        # Extract pages
        if hasattr(result, "pages"):
            for page in result.pages:
                page_data = {
                    "page_number": page.page_number,
                    "width": page.width if hasattr(page, "width") else None,
                    "height": page.height if hasattr(page, "height") else None,
                    "unit": page.unit if hasattr(page, "unit") else None,
                    "lines": [],
                }

                # Extract lines (text content)
                if hasattr(page, "lines"):
                    for line in page.lines:
                        page_data["lines"].append(
                            {
                                "content": line.content,
                                "polygon": line.polygon if hasattr(line, "polygon") else None,
                            }
                        )

                structured["pages"].append(page_data)

        # Extract tables
        if hasattr(result, "tables"):
            for table in result.tables:
                table_data = {
                    "row_count": table.row_count if hasattr(table, "row_count") else 0,
                    "column_count": table.column_count if hasattr(table, "column_count") else 0,
                    "cells": [],
                }

                if hasattr(table, "cells"):
                    for cell in table.cells:
                        table_data["cells"].append(
                            {
                                "row_index": cell.row_index if hasattr(cell, "row_index") else None,
                                "column_index": cell.column_index
                                if hasattr(cell, "column_index")
                                else None,
                                "content": cell.content,
                                "kind": cell.kind if hasattr(cell, "kind") else None,
                            }
                        )

                structured["tables"].append(table_data)

        # Extract paragraphs
        if hasattr(result, "paragraphs"):
            for paragraph in result.paragraphs:
                structured["paragraphs"].append(
                    {
                        "content": paragraph.content,
                        "role": paragraph.role if hasattr(paragraph, "role") else None,
                    }
                )

        return structured

    async def _parse_excel(self, document_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Excel file using openpyxl.

        Args:
            document_bytes: Excel binary content
            filename: Original filename

        Returns:
            Structured Excel data with sheets and cells

        """
        from io import BytesIO

        import openpyxl

        workbook = openpyxl.load_workbook(BytesIO(document_bytes), data_only=True)

        structured = {
            "filename": filename,
            "sheets": [],
            "metadata": {"format": "excel", "sheet_count": len(workbook.sheetnames)},
        }

        for sheet in workbook.worksheets:
            sheet_data = {
                "name": sheet.title,
                "row_count": sheet.max_row,
                "column_count": sheet.max_column,
                "rows": [],
            }

            for row in sheet.iter_rows(values_only=True):
                sheet_data["rows"].append([cell for cell in row])

            structured["sheets"].append(sheet_data)

        logger.info(
            f"Parsed Excel file: {filename} - {len(structured['sheets'])} sheets, "
            f"{structured['sheets'][0]['row_count'] if structured['sheets'] else 0} rows"
        )

        return structured

    async def _parse_word(self, document_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Word document using python-docx.

        Args:
            document_bytes: Word binary content
            filename: Original filename

        Returns:
            Structured Word data with paragraphs and tables

        """
        from io import BytesIO

        import docx

        document = docx.Document(BytesIO(document_bytes))

        structured = {
            "filename": filename,
            "paragraphs": [],
            "tables": [],
            "metadata": {"format": "word"},
        }

        # Extract paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                structured["paragraphs"].append(
                    {
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else None,
                    }
                )

        # Extract tables
        for idx, table in enumerate(document.tables):
            table_data = {
                "table_number": idx + 1,
                "row_count": len(table.rows),
                "column_count": len(table.columns) if table.rows else 0,
                "cells": [],
            }
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["cells"].append(row_data)
            structured["tables"].append(table_data)

        logger.info(
            f"Parsed Word file: {filename} - {len(structured['paragraphs'])} paragraphs, "
            f"{len(structured['tables'])} tables"
        )

        return structured

    async def _handle_parsing_failure(
        self, blob_path: str, filename: str, exception: Exception
    ) -> None:
        """
        Handle parsing failure by moving document to dead letter queue.

        Args:
            blob_path: Original blob path (format: "container/blob_name")
            filename: Document filename
            exception: Exception that caused failure
        """
        try:
            # Parse blob path
            parts = blob_path.split("/", 1)
            if len(parts) != 2:
                logger.error(f"Invalid blob path format: {blob_path}")
                return

            source_container, source_blob = parts

            # Build error details (all fields will be preserved in DLQ metadata)
            error_details = {
                "error_type": exception.__class__.__name__,
                "error_message": str(exception),
                "timestamp": datetime.utcnow().isoformat(),
                "filename": filename,
                "operation": "document_parsing",
            }

            # Move to DLQ (await blob client getter)
            blob_client = await self._get_blob_client()
            dlq_path = await blob_client.move_to_dead_letter_queue(
                source_container=source_container,
                source_blob=source_blob,
                error_details=error_details,
            )

            logger.info(f"Moved failed document to DLQ: {blob_path} → {dlq_path}")

        except Exception as dlq_exc:
            logger.error(f"Failed to move document to DLQ: {dlq_exc}")

    async def close(self) -> None:
        """
        Close clients and release resources.

        Should be called during application shutdown.
        """
        if self._di_client is not None:
            await self._di_client.close()
            logger.info("Closed DocumentIntelligenceClient")

        if self._blob_client is not None:
            await self._blob_client.close()
